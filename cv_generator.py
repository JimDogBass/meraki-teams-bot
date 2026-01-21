"""
CV Generator - Creates Meraki-formatted Word documents from structured CV data.
Logo in header (centered) so it appears on every page.
"""
import io
import json
import os
import copy
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT

# Logo path
LOGO_PATH = os.path.join(os.path.dirname(__file__), "templates", "meraki_logo.png")

# Tab positions for two-column layout (matching Emeliene CV)
TAB_RIGHT_POS = Cm(3.5)  # Right-aligned tab for labels
TAB_LEFT_POS = Cm(4.5)   # Left-aligned tab for values

# Grey shading for section headers
HEADER_SHADING = "D9D9D9"


def add_page_border(doc):
    """Add a single-line border around the page."""
    for section in doc.sections:
        sectPr = section._sectPr
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')

        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '24')
            border.set(qn('w:color'), 'auto')
            pgBorders.append(border)

        sectPr.append(pgBorders)


def add_header_logo(doc, logo_path):
    """
    Add logo to header (centered, appears on every page).
    Note: Headers appear slightly faded in Word's editing view but print full color.
    """
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False

        # Get or create paragraph in header
        if not header.paragraphs:
            header.add_paragraph()
        para = header.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add centered logo
        run = para.add_run()
        run.add_picture(logo_path, width=Inches(2.5))


def create_meraki_cv(cv_data: dict) -> bytes:
    """
    Generate a Meraki-formatted CV Word document from structured data.
    Builds document from scratch with logo in body (not header).
    """
    # Create new blank document
    doc = Document()

    # Set default font to Aptos 11pt with no spacing after paragraphs
    style = doc.styles['Normal']
    style.font.name = 'Aptos'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # Set page margins (reduced top margin for logo)
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Add page border
    add_page_border(doc)

    # === ADD LOGO IN BODY (centered, page 1 only, full color) ===
    if os.path.exists(LOGO_PATH):
        add_blank_line(doc)  # Space before logo
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_para.add_run()
        logo_run.add_picture(LOGO_PATH, width=Inches(3.5))
        add_blank_line(doc)  # Space after logo

    # === PERSONAL DETAILS ===
    add_section_header(doc, "PERSONAL DETAILS")
    add_blank_line(doc)
    add_field_line(doc, "Name", cv_data.get("name", ""))
    # Only show these fields if they have values
    if cv_data.get("location"):
        add_field_line(doc, "Location", cv_data.get("location", ""))
    if cv_data.get("right_to_work"):
        add_field_line(doc, "Right to Work", cv_data.get("right_to_work", ""))
    if cv_data.get("notice"):
        add_field_line(doc, "Notice", cv_data.get("notice", ""))
    if cv_data.get("salary_expectations"):
        add_field_line(doc, "Salary expectations", cv_data.get("salary_expectations", ""))

    # === IT/SYSTEMS (always show, even if empty) ===
    add_blank_line(doc)
    it_systems = cv_data.get("it_systems", "")
    p = doc.add_paragraph()
    label_run = p.add_run("IT/Systems: ")
    label_run.bold = True
    p.add_run(it_systems if it_systems else "N/A")

    # === LANGUAGES (always show, even if empty) ===
    languages = cv_data.get("languages", "")
    p = doc.add_paragraph()
    label_run = p.add_run("Languages: ")
    label_run.bold = True
    p.add_run(languages if languages else "N/A")

    # === INTERESTS (always show, even if empty) ===
    interests = cv_data.get("interests", "")
    p = doc.add_paragraph()
    label_run = p.add_run("Interests: ")
    label_run.bold = True
    p.add_run(interests if interests else "N/A")

    # === EDUCATION ===
    education = cv_data.get("education", [])
    professional_qualifications = cv_data.get("professional_qualifications", [])

    if education or professional_qualifications:
        add_blank_line(doc)
        add_section_header(doc, "EDUCATION")
        add_blank_line(doc)

        # Professional Qualifications first (if any)
        if professional_qualifications:
            p = doc.add_paragraph()
            label_run = p.add_run("Professional Qualifications:")
            label_run.bold = True
            for qual in professional_qualifications:
                # Use en dash bullet for professional qualifications
                qual_p = doc.add_paragraph()
                qual_p.paragraph_format.left_indent = Cm(0.6)
                qual_p.paragraph_format.first_line_indent = Cm(-0.4)
                qual_p.add_run(f"−\t{qual}")
            add_blank_line(doc)

        for i, edu in enumerate(education):
            dates = edu.get("dates", "")
            institution = edu.get("institution", "")
            title = edu.get("title", "")
            details = edu.get("details", [])

            # Date and qualification line (bold)
            add_tabbed_line(doc, dates, title, bold=True)

            # Institution on next line (indented)
            if institution:
                add_indented_line(doc, institution)

            # Additional details
            for detail in details:
                add_indented_line(doc, detail)

            # Add blank line between education entries (except last)
            if i < len(education) - 1:
                add_blank_line(doc)

    # === CANDIDATE PROFILE (if present) ===
    profile = cv_data.get("profile", "")
    if profile:
        add_blank_line(doc)
        add_section_header(doc, "CANDIDATE PROFILE")
        add_blank_line(doc)
        # Split profile into paragraphs if it contains double newlines
        paragraphs = profile.split('\n\n') if '\n\n' in profile else [profile]
        for para_text in paragraphs:
            p = doc.add_paragraph(para_text.strip())
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # === WORK EXPERIENCE ===
    work_exp = cv_data.get("work_experience", [])
    if work_exp:
        add_blank_line(doc)
        add_blank_line(doc)
        add_section_header(doc, "WORK EXPERIENCE")
        add_blank_line(doc)
        for i, job in enumerate(work_exp):
            dates = job.get("dates", "")
            company = job.get("company", "")
            position = job.get("position", "")

            # Date and company line (bold)
            add_tabbed_line(doc, dates, company, bold=True)

            # Position line
            if position:
                add_position_line(doc, position)

            # Work experience sections (sub-headers with bullet points)
            sections = job.get("sections", [])
            if sections:
                for section in sections:
                    header = section.get("header", "")
                    content = section.get("content", [])

                    # Add sub-header if present (bold text, no bullet)
                    if header:
                        add_blank_line(doc)
                        add_work_section_header(doc, header)

                    # Add content as bullet points (supports nested structure)
                    for item in content:
                        add_nested_bullets(doc, item)
            else:
                # Fallback: handle old "bullets" format for backwards compatibility
                for bullet in job.get("bullets", []):
                    add_nested_bullets(doc, bullet)

            # Add blank line between jobs (except last)
            if i < len(work_exp) - 1:
                add_blank_line(doc)

    # === OTHER INFORMATION (only if present and has content) ===
    other_info = cv_data.get("other_information", [])
    if other_info and len(other_info) > 0:
        add_blank_line(doc)
        add_section_header(doc, "OTHER INFORMATION")
        add_blank_line(doc)
        for item in other_info:
            category = item.get("category", "")
            content = item.get("content", [])
            entries = item.get("entries", [])

            if entries:
                # Detailed entries (e.g., Non-Profit Boards, Volunteer Work)
                # Add category header
                p = doc.add_paragraph()
                label_run = p.add_run(category)
                label_run.bold = True
                label_run.underline = True

                for entry in entries:
                    org = entry.get("organization", "")
                    dates = entry.get("dates", "")
                    role = entry.get("role", "")
                    bullets = entry.get("bullets", [])

                    # Organization and dates line (bold)
                    if org or dates:
                        add_tabbed_line(doc, dates, org, bold=True)

                    # Role on next line (indented)
                    if role:
                        add_indented_line(doc, role)

                    # Bullet points
                    for bullet in bullets:
                        add_nested_bullets(doc, bullet)

                add_blank_line(doc)
            elif content:
                # Simple content - category header then each item on its own line
                p = doc.add_paragraph()
                if category:
                    label_run = p.add_run(category + ":")
                    label_run.bold = True

                # Each content item as a separate line
                if isinstance(content, list) and len(content) > 0:
                    for content_item in content:
                        item_p = doc.add_paragraph(content_item)
                else:
                    p.add_run(str(content))

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def add_blank_line(doc):
    """Add an empty paragraph for spacing."""
    return doc.add_paragraph()


def add_section_header(doc, text):
    """Add a bold section header with grey shading."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Aptos'
    run.font.size = Pt(11)

    # Add grey shading to paragraph
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), HEADER_SHADING)
    pPr.append(shd)

    return p


def add_field_line(doc, label, value):
    """Add a field line with bold tabbed label and value (e.g., '\tName\tJohn Smith')."""
    p = doc.add_paragraph()
    # Set up tab stops
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(TAB_RIGHT_POS, WD_TAB_ALIGNMENT.RIGHT)
    tab_stops.add_tab_stop(TAB_LEFT_POS, WD_TAB_ALIGNMENT.LEFT)
    # Add tabbed content with bold label
    p.add_run("\t")
    label_run = p.add_run(label)
    label_run.bold = True
    p.add_run(f"\t{value if value else ''}")
    return p


def add_tabbed_line(doc, left_text, right_text, bold=False):
    """Add a line with tab-separated content (e.g., '\t2019\tMA Event Design')."""
    p = doc.add_paragraph()
    # Set up tab stops
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(TAB_RIGHT_POS, WD_TAB_ALIGNMENT.RIGHT)
    tab_stops.add_tab_stop(TAB_LEFT_POS, WD_TAB_ALIGNMENT.LEFT)
    # Set hanging indent so wrapped text aligns with right column
    p.paragraph_format.left_indent = TAB_LEFT_POS
    p.paragraph_format.first_line_indent = -TAB_LEFT_POS
    # Add tabbed content with optional bold
    run = p.add_run(f"\t{left_text}\t{right_text}")
    if bold:
        run.bold = True
    return p


def add_indented_line(doc, text):
    """Add an indented line (for institution names, details)."""
    p = doc.add_paragraph()
    # Set up BOTH tab stops (same as main lines) for proper alignment
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(TAB_RIGHT_POS, WD_TAB_ALIGNMENT.RIGHT)
    tab_stops.add_tab_stop(TAB_LEFT_POS, WD_TAB_ALIGNMENT.LEFT)
    # Set hanging indent so wrapped text aligns with content column
    p.paragraph_format.left_indent = TAB_LEFT_POS
    p.paragraph_format.first_line_indent = -TAB_LEFT_POS
    # Add double-tabbed content (first tab to RIGHT pos, second to LEFT pos)
    p.add_run(f"\t\t{text}")
    return p


def add_position_line(doc, position):
    """Add a position line with bold label and bold position value."""
    p = doc.add_paragraph()
    # Set up tab stops
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(TAB_RIGHT_POS, WD_TAB_ALIGNMENT.RIGHT)
    tab_stops.add_tab_stop(TAB_LEFT_POS, WD_TAB_ALIGNMENT.LEFT)
    # Set hanging indent so wrapped lines align with position value
    p.paragraph_format.left_indent = TAB_LEFT_POS
    p.paragraph_format.first_line_indent = -TAB_LEFT_POS
    # Add tabbed position (both label and value bold)
    p.add_run("\t")
    label_run = p.add_run("Position:")
    label_run.bold = True
    p.add_run("\t")
    value_run = p.add_run(position)
    value_run.bold = True
    return p


def add_work_section_header(doc, header_text):
    """Add a bold sub-header within work experience (e.g., 'Client Product Strategy & Bespoke Benchmark Design')."""
    p = doc.add_paragraph()
    run = p.add_run(header_text)
    run.bold = True
    run.font.name = 'Aptos'
    run.font.size = Pt(11)
    return p


def add_work_content_paragraph(doc, text):
    """Add a content paragraph within work experience (plain text, justified)."""
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_bullet_point(doc, text, level=0):
    """Add a bullet point line with proper hanging indent. Supports 4 levels."""
    p = doc.add_paragraph()

    # Different bullets for 4 levels (matching typical CV hierarchy)
    # Level 0: filled bullet, Level 1: en dash, Level 2: arrow, Level 3: open circle
    bullets = ["•", "–", "›", "○"]
    bullet_char = bullets[min(level, len(bullets) - 1)]

    # Indentation increases with level
    base_indent = Cm(0.6)
    level_indent = Cm(0.5 * level)

    p.paragraph_format.left_indent = base_indent + level_indent
    p.paragraph_format.first_line_indent = Cm(-0.4)
    p.add_run(f"{bullet_char}\t{text}")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_nested_bullets(doc, bullet_item, level=0):
    """
    Recursively add bullets with nested structure.
    bullet_item can be:
    - A string (simple bullet)
    - A dict with 'text' and optional 'sub_bullets'
    """
    if isinstance(bullet_item, str):
        # Simple string bullet
        add_bullet_point(doc, bullet_item, level)
    elif isinstance(bullet_item, dict):
        # Nested bullet with potential sub-bullets
        text = bullet_item.get("text", "")
        if text:
            add_bullet_point(doc, text, level)

        # Process sub-bullets recursively
        sub_bullets = bullet_item.get("sub_bullets", [])
        for sub_item in sub_bullets:
            add_nested_bullets(doc, sub_item, level + 1)


def parse_cv_json(ai_response: str) -> dict:
    """
    Parse the AI response to extract JSON CV data.
    Handles cases where JSON might be wrapped in markdown code blocks.
    """
    text = ai_response.strip()

    # Remove markdown code blocks if present
    if text.startswith("```json"):
        text = text[7:]
    elif text.startswith("```"):
        text = text[3:]

    if text.endswith("```"):
        text = text[:-3]

    text = text.strip()

    try:
        return json.loads(text)
    except json.JSONDecodeError as e:
        # Try to find JSON object in the response
        start = text.find('{')
        end = text.rfind('}')
        if start != -1 and end != -1:
            try:
                return json.loads(text[start:end+1])
            except json.JSONDecodeError:
                pass
        raise ValueError(f"Could not parse CV data as JSON: {e}")


# System prompt for CV extraction - preserves full detail
CV_EXTRACTION_PROMPT = """You are a CV data extraction assistant. Extract structured information from the provided CV and return it as valid JSON.

CRITICAL RULES:
1. Return ONLY valid JSON - no explanations, no markdown code blocks, just the JSON object
2. Use short date format: "Jan 23" not "January 2023"
3. Extract ALL work experience entries
4. Preserve ALL content from each role
5. Extract IT/Systems separately - do NOT include them in other_information
6. Only include "other_information" for categories not covered by other fields
7. NEVER FABRICATE OR HALLUCINATE CONTENT - only extract what is EXPLICITLY written in the CV. If information is not present, use empty string "" or empty array []. Do NOT invent volunteering, skills, experiences, or any other content.
8. NEVER use "N/A" for any field - always use empty string "" instead

REQUIRED JSON STRUCTURE:
{
  "name": "Full Name",
  "location": "City, Country",
  "right_to_work": "",
  "notice": "",
  "salary_expectations": "",
  "it_systems": "Comma-separated list of software, systems, tools (e.g., Salesforce, Backstop, Dealcloud, Excel, Bloomberg)",
  "qualifications": "",
  "languages": "Comma-separated list of languages with proficiency (e.g., English (native), French (proficient), Spanish (conversational))",
  "interests": "Comma-separated list of interests/hobbies (e.g., Travel, volunteering, contemporary art, basketball)",
  "profile": "The candidate's personal profile/summary/about section - a dedicated paragraph about their career and professional background. NOT hobbies, NOT interests, NOT education descriptions, NOT key skills lists. Must be from a section explicitly labeled Profile/Summary/About. If no such section exists, use empty string.",
  "professional_qualifications": [
    "CFA Level 2 Candidate (2026)",
    "Quantamental Academy – Macrosynergy (2025)",
    "Portfolio Risk Management (HEC Paris @ Coursera, 2023)"
  ],
  "education": [
    {
      "dates": "2019",
      "title": "Degree/Qualification name",
      "institution": "University/School name",
      "details": ["Additional details like grades, citations, leadership roles, honors if present"]
    }
  ],
  "work_experience": [
    {
      "dates": "Jan 23 - Present",
      "company": "Company Name",
      "position": "Job Title",
      "sections": [
        {
          "header": "Client Product Strategy & Bespoke Benchmark Design",
          "content": [
            "Acted as the lead product advisor for institutional clients...",
            "Developed compelling index investment narratives..."
          ]
        },
        {
          "header": "Systematic Research & Portfolio Optimization",
          "content": [
            "Designed and implemented a Python-based backtesting framework..."
          ]
        }
      ]
    }
  ],
  "other_information": [
    {
      "category": "Core Skills",
      "content": ["Skill 1", "Skill 2", "Skill 3"]
    },
    {
      "category": "Volunteering",
      "content": [
        "Volunteering for Food for all to deliver meals to the vulnerable",
        "Volunteering for the LTA - umpiring tennis matches for the visually impaired"
      ]
    }
  ]
}

CRITICAL - TITLE PROGRESSIONS AT THE SAME COMPANY:
- This rule ONLY applies when the EXACT SAME company name appears with multiple titles listed together
- DIFFERENT companies MUST be separate work_experience entries - NEVER merge different companies together
- UBS, Credit Suisse, Goldman Sachs, etc. are DIFFERENT companies even if the person moved between them
- Example of what to MERGE (same company, multiple titles listed together):
    "Lincolnshire Management
     Private Equity Associate                    September 2024 – Present
     Private Equity Analyst, Senior Analyst      February 2022 – September 2024"
  → ONE entry: dates "Feb 22 - Present", position shows all titles
- Example of what to KEEP SEPARATE (different companies):
    "UBS - Execution Only Trader (June 2025 - Present)
     Credit Suisse - Multi-Asset Trader (May 2023 - June 2025)
     Louis Capital Markets - Equity Trader (May 2018 - May 2023)"
  → THREE separate entries, one for each company
- When in doubt, keep jobs as SEPARATE entries

IMPORTANT - WORK EXPERIENCE SECTIONS:
- Many CVs organize work experience with SUB-HEADERS followed by bullet points
- Sub-headers are STANDALONE section titles like "Client Product Strategy & Bespoke Benchmark Design" that appear on their OWN line with bullet points BELOW them
- Use the "sections" array ONLY when there are clear standalone section headers within a role
- If a role has NO sub-headers and just flat bullet points, use a single section with empty header: {"header": "", "content": ["bullet 1", "bullet 2"]}

CRITICAL - DO NOT SPLIT BULLET POINTS:
- If a bullet point has "Label: content" format like "Strategic Leadership: Define and implement the overall business strategy...", this is ONE bullet point - keep it as ONE bullet point
- DO NOT extract "Strategic Leadership" as a header and "Define and implement..." as a separate bullet
- The ENTIRE text "Strategic Leadership: Define and implement the overall business strategy..." should be a SINGLE bullet point
- Only use headers when they are STANDALONE lines in the original CV, not when they're part of a bullet point

- PRESERVE the exact text of each content item
- Content items can be simple strings OR nested objects with sub_bullets for hierarchical bullets:
  - Simple: "Strategic Leadership: Define and implement the overall business strategy..."
  - Nested: {"text": "Main point", "sub_bullets": ["Sub point 1", "Sub point 2"]}

CRITICAL - SUMMARY-STYLE CVs:
- Some CVs list ALL roles first (just dates, company, title) and THEN have a "Professional Highlights" or "Key Responsibilities" section with bullet points that describe experience ACROSS ALL roles
- DO NOT put these summary bullets under just the first/current role!
- If bullet points are clearly a SUMMARY of career experience (not specific to one role), distribute them appropriately:
  - If bullets mention specific companies/roles, assign them to those roles
  - If bullets are generic highlights that apply across roles, assign them to the MOST RECENT role that they're relevant to based on context
  - Look for clues: "Worked with DWS in New York" → assign to the New York role; "Managed portfolio of internal audit" → assign based on where that work happened
- Each role should have its OWN relevant bullets, not one role with everything and others empty
- The "profile" field should contain the candidate's personal statement/summary paragraph, NOT the professional highlights bullets

PROFESSIONAL QUALIFICATIONS:
- "professional_qualifications" is an ARRAY of strings for ALL certificates, courses, learning, professional development
- This includes: CFA, ACCA, ACA, CAIA, PRINCE2, Scrum Master, PMP, Six Sigma, Coursera courses, company training, conference attendance, seminars, academies, etc.
- Look for sections titled "Certificates", "Certifications", "Learning", "Professional Development", "Outreach", "Training", "Courses and Certificates", "Additional Information"
- Extract ALL courses and certificates regardless of how old they are
- Each item should be a complete string like "CFA Level 2 Candidate (2026)" or "Nexia International – European Manager training seminar (2007)"
- If no professional qualifications found, set to empty array []

OTHER RULES:
- ABSOLUTELY NO HALLUCINATION: Only extract content that is EXPLICITLY written in the CV. Do NOT invent, fabricate, or assume ANY content. If you are unsure whether something exists in the CV, it probably doesn't - leave it out.
- Extract ALL paid/professional roles from the CV as work_experience
- Non-Profit Boards, Volunteer roles, and Advisory roles should go in "other_information" with full details preserved - but ONLY if they actually exist in the CV
- Preserve original wording exactly - do NOT paraphrase or add content
- "profile" should contain the candidate's PERSONAL STATEMENT or ABOUT ME section - typically a dedicated section labeled "Profile", "Summary", "About Me", or "Personal Statement". This is NOT: hobbies/interests text, key skills lists, education study descriptions, course descriptions, or bullet point achievements. Random sentences from hobbies (like "I am a keen golfer") are NOT profiles. If no dedicated profile/summary section exists, set to empty string ""
- NEVER use "N/A" for any field - use empty string "" instead
- "location" should only contain city/country if explicitly stated. If not found, set to empty string ""
- NEVER infer right_to_work from nationality, citizenship, visa status, or "eligible to work" statements. The right_to_work field must ONLY be filled if the CV has an EXPLICIT "Right to Work:" label. Otherwise set to empty string ""
- NEVER infer or guess notice or salary_expectations - ALWAYS leave these as empty strings "" unless EXPLICITLY stated with those exact labels in the CV
- "it_systems" should contain ANY software, systems, or tools mentioned in the CV (CRM, databases, financial platforms, Microsoft Office, etc.) - extract from skills sections, bullet points, or anywhere mentioned. If no IT/systems found, set to empty string ""
- "qualifications" field is deprecated - always set to empty string ""
- "languages" should contain ALL languages mentioned with proficiency level in parentheses (e.g., "English (native), French (proficient)"). If no languages found, set to empty string ""
- "interests" should contain hobbies/interests as a SHORT comma-separated list of categories (e.g., "Travel, Sports, Music"). If the CV has DETAILED hobby descriptions, put those full descriptions in "other_information" under a "Hobbies" or "Interests" category so they're preserved
- Education "details" should include ALL additional info: grades, GPA, honors, citations (e.g., "Citation in French"), leadership roles (e.g., "Leadership: President, Harvard College European Society"), test scores, study profile descriptions (e.g., "Profile: Marketing and Management"), faculty info, thesis titles, etc. These belong in education details, NOT in the profile field.
- "other_information" is a CATCH-ALL for anything not covered by other fields. BE INCLUSIVE - if in doubt, include it here. Examples: Publications, Awards, Core Skills, Key Skills, Volunteer Work, detailed hobby descriptions, summary of experience bullets, achievements, anything else the candidate included. Use "content": ["item1", "item2"] format with the ORIGINAL text preserved exactly. Better to include too much than lose information - the recruiter can delete what they don't need. If nothing else remains, set to empty array []

Extract the CV data now:"""
