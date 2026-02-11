"""
Fernando Format - CV Reformat Bot for Meraki Talent
Single-function bot that reformats CVs with Meraki branding and generates alternative candidate profiles.
Stripped down from "Jimmy Content" bot to focus solely on CV reformatting.
"""
import nest_asyncio
nest_asyncio.apply()

import os
import asyncio
import io
import time
import httpx
import tempfile
import subprocess
import base64
import urllib.parse
from flask import Flask, request, Response
from botbuilder.core import BotFrameworkAdapter, BotFrameworkAdapterSettings, TurnContext
from botbuilder.schema import Activity, Attachment
from openai import AzureOpenAI
import pdfplumber
import pytesseract
from pdf2image import convert_from_bytes
from docx import Document
from azure.data.tables import TableServiceClient
from azure.storage.blob import BlobServiceClient, generate_blob_sas, BlobSasPermissions
from datetime import datetime, timedelta
from cv_generator import create_meraki_cv, parse_cv_json, CV_EXTRACTION_PROMPT
import re
from html import unescape
import msal

app = Flask(__name__)

# Bot Framework adapter setup
settings = BotFrameworkAdapterSettings(
    app_id=os.environ.get("MICROSOFT_APP_ID", ""),
    app_password=os.environ.get("MICROSOFT_APP_PASSWORD", ""),
    channel_auth_tenant=os.environ.get("MICROSOFT_APP_TENANT_ID", "")
)
adapter = BotFrameworkAdapter(settings)

# Azure OpenAI client setup
openai_client = AzureOpenAI(
    api_key=os.environ.get("AZURE_OPENAI_KEY"),
    api_version="2024-02-01",
    azure_endpoint=os.environ.get("AZURE_OPENAI_ENDPOINT")
)

# Azure Storage connection
storage_connection_string = os.environ.get("AZURE_STORAGE_CONNECTION_STRING", "")
table_service_client = None
blob_service_client = None
if storage_connection_string:
    table_service_client = TableServiceClient.from_connection_string(storage_connection_string)
    blob_service_client = BlobServiceClient.from_connection_string(storage_connection_string)
    try:
        blob_service_client.create_container("cv-outputs")
    except Exception:
        pass

# Pending state TTL - for button click -> file upload flow
PENDING_ROLE_TTL_SECONDS = 300  # 5 minutes

# Alternative Candidate Profile prompt
ALTERNATIVE_PROFILE_PROMPT = """Based on this CV data, write a short alternative candidate profile (2-3 sentences max).

This is a punchy summary to send alongside the CV to a client. Use "they/their" pronouns. Highlight their key strengths, current situation, and what makes them stand out. Keep it compelling and concise.

CV Data:
{cv_json}

Output ONLY the profile text, nothing else."""


def get_pending_reformat(conversation_id: str) -> bool:
    """Check if there's a pending reformat request for this conversation."""
    if not table_service_client:
        return False

    try:
        table_client = table_service_client.get_table_client("BotState")
        row_key = str(hash(conversation_id) & 0xFFFFFFFF)

        try:
            entity = table_client.get_entity(partition_key="pending_role", row_key=row_key)
            timestamp = entity.get("Timestamp_", 0)

            if time.time() - timestamp < PENDING_ROLE_TTL_SECONDS:
                return True
            else:
                table_client.delete_entity(partition_key="pending_role", row_key=row_key)
        except Exception:
            pass
    except Exception as e:
        print(f"[ERROR] Getting pending state: {e}")

    return False


def set_pending_reformat(conversation_id: str):
    """Set a pending reformat request in Azure Table Storage."""
    if not table_service_client:
        return

    try:
        table_client = table_service_client.get_table_client("BotState")

        try:
            table_service_client.create_table("BotState")
        except Exception:
            pass

        row_key = str(hash(conversation_id) & 0xFFFFFFFF)
        entity = {
            "PartitionKey": "pending_role",
            "RowKey": row_key,
            "ConversationId": conversation_id[:900],
            "Role": "reformat",
            "Timestamp_": time.time()
        }
        table_client.upsert_entity(entity)
    except Exception as e:
        print(f"[ERROR] Setting pending state: {e}")


def clear_pending_reformat(conversation_id: str):
    """Clear pending reformat state from Azure Table Storage."""
    if not table_service_client:
        return

    try:
        table_client = table_service_client.get_table_client("BotState")
        row_key = str(hash(conversation_id) & 0xFFFFFFFF)
        table_client.delete_entity(partition_key="pending_role", row_key=row_key)
    except Exception:
        pass


def extract_text_from_html(html_content: str) -> str:
    """Extract plain text from HTML content."""
    if not html_content:
        return ""
    text = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'</p>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'</li>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'</tr>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'</td>', ' | ', text, flags=re.IGNORECASE)
    text = re.sub(r'<[^>]+>', '', text)
    text = unescape(text)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n\s*\n', '\n\n', text)
    text = text.strip()
    return text


def is_cv_content(text: str) -> bool:
    """Check if text looks like CV content."""
    if len(text) < 200:
        return False
    cv_indicators = ['experience', 'education', 'skills', 'accomplishment', 'employment',
                     'professional', 'qualification', 'certification', 'university', 'degree']
    text_lower = text.lower()
    matches = sum(1 for indicator in cv_indicators if indicator in text_lower)
    return matches >= 2


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF file bytes, including tables. Falls back to OCR for image-based PDFs."""
    text_parts = []

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    row_text = [cell.strip() if cell else "" for cell in row]
                    row_text = [cell for cell in row_text if cell]
                    if row_text:
                        text_parts.append(" | ".join(row_text))

    text = "\n".join(text_parts).strip()

    # If pdfplumber found no text, try OCR
    if not text:
        try:
            images = convert_from_bytes(file_bytes)
            ocr_parts = []
            for image in images:
                page_text = pytesseract.image_to_string(image)
                if page_text.strip():
                    ocr_parts.append(page_text.strip())
            text = "\n".join(ocr_parts)
        except Exception as e:
            print(f"[ERROR] OCR failed: {e}")

    return text


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from Word document bytes, including nested tables."""
    doc = Document(io.BytesIO(file_bytes))

    text_parts = [para.text for para in doc.paragraphs]

    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                text_parts.append(" | ".join(row_text))

    standard_text = "\n".join(text_parts).strip()

    # If standard extraction got very little text, use deep XML extraction
    if len(standard_text) < 500:
        try:
            body = doc._body._body
            texts = []
            for child in body.iter():
                if child.tag.endswith('}t'):
                    if child.text:
                        texts.append(child.text)
            deep_text = ''.join(texts)
            if len(deep_text) > len(standard_text):
                return deep_text
        except Exception:
            pass

    return standard_text


def extract_text_from_doc(file_bytes: bytes) -> str:
    """Extract text from old .doc format using antiword."""
    with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
        tmp.write(file_bytes)
        tmp.flush()
        tmp_path = tmp.name

    try:
        result = subprocess.run(
            ['antiword', tmp_path],
            capture_output=True,
            text=True,
            timeout=30
        )
        if result.returncode == 0:
            return result.stdout.strip()
        else:
            raise ValueError(f"antiword failed: {result.stderr}")
    finally:
        os.unlink(tmp_path)


def get_graph_token() -> str:
    """Get Microsoft Graph API token using app credentials."""
    app_id = os.environ.get("MICROSOFT_APP_ID", "")
    app_password = os.environ.get("MICROSOFT_APP_PASSWORD", "")
    tenant_id = os.environ.get("MICROSOFT_APP_TENANT_ID", "")

    if not all([app_id, app_password, tenant_id]):
        return None

    authority = f"https://login.microsoftonline.com/{tenant_id}"
    msal_app = msal.ConfidentialClientApplication(
        app_id,
        authority=authority,
        client_credential=app_password
    )

    result = msal_app.acquire_token_for_client(scopes=["https://graph.microsoft.com/.default"])
    if "access_token" in result:
        return result["access_token"]
    else:
        print(f"[ERROR] Failed to get Graph token: {result.get('error_description', result)}")
        return None


async def get_files_from_chat(chat_id: str, token: str, user_aad_id: str = None) -> list:
    """Fetch recent files from a Teams chat using Graph API."""
    files = []
    headers = {"Authorization": f"Bearer {token}"}

    graph_chat_id = chat_id
    if graph_chat_id.startswith('a:'):
        graph_chat_id = graph_chat_id[2:]

    # Try to find chat by user if we have their AAD ID
    actual_chat_id = None
    if user_aad_id:
        try:
            async with httpx.AsyncClient() as client:
                chats_url = f"https://graph.microsoft.com/v1.0/users/{user_aad_id}/chats?$top=20"
                resp = await client.get(chats_url, headers=headers)
                if resp.status_code == 200:
                    chats_data = resp.json()
                    for chat in chats_data.get("value", []):
                        if chat.get("chatType") == "oneOnOne":
                            actual_chat_id = chat.get("id")
                            break
        except Exception as e:
            print(f"[ERROR] Failed to list user chats: {e}")

    if not actual_chat_id:
        actual_chat_id = urllib.parse.quote(graph_chat_id, safe='')

    url = f"https://graph.microsoft.com/v1.0/chats/{actual_chat_id}/messages?$top=10"

    async with httpx.AsyncClient() as client:
        response = await client.get(url, headers=headers)

        if response.status_code != 200:
            return files

        data = response.json()
        for message in data.get("value", []):
            attachments = message.get("attachments", [])
            for att in attachments:
                name = att.get("name", "")
                if att.get("contentUrl"):
                    files.append({
                        "name": name,
                        "content_url": att.get("contentUrl"),
                        "content_type": att.get("contentType", "")
                    })

    return files


async def download_file_from_sharepoint(file_url: str, token: str) -> bytes:
    """Download a file from SharePoint using Graph API."""
    headers = {"Authorization": f"Bearer {token}"}

    async with httpx.AsyncClient() as client:
        if "sharepoint.com" in file_url:
            encoded_url = base64.urlsafe_b64encode(file_url.encode()).decode().rstrip('=')
            share_id = f"u!{encoded_url}"
            graph_url = f"https://graph.microsoft.com/v1.0/shares/{share_id}/driveItem/content"
            response = await client.get(graph_url, headers=headers, follow_redirects=True)
        else:
            response = await client.get(file_url, headers=headers, follow_redirects=True)

        response.raise_for_status()
        return response.content


async def download_attachment(attachment, turn_context: TurnContext) -> bytes:
    """Download attachment from Teams."""
    download_url = None

    if attachment.content_type == "application/vnd.microsoft.teams.file.download.info":
        if isinstance(attachment.content, dict) and "downloadUrl" in attachment.content:
            download_url = attachment.content["downloadUrl"]

    if not download_url and attachment.content_url:
        download_url = attachment.content_url

    if not download_url and isinstance(attachment.content, dict):
        download_url = attachment.content.get("downloadUrl") or attachment.content.get("download_url")

    if not download_url:
        raise ValueError(f"No download URL found in attachment")

    async with httpx.AsyncClient() as client:
        response = await client.get(download_url, follow_redirects=True)
        response.raise_for_status()
        return response.content


async def extract_text_from_attachment(attachment, turn_context: TurnContext) -> str:
    """Download and extract text from an attachment."""
    name = attachment.name or ""
    if not name and isinstance(attachment.content, dict):
        name = attachment.content.get("name", "")
    content_type = attachment.content_type or ""

    try:
        file_bytes = await download_attachment(attachment, turn_context)

        if name.lower().endswith('.pdf') or 'pdf' in content_type.lower():
            return extract_text_from_pdf(file_bytes)
        elif name.lower().endswith('.docx') or 'wordprocessingml' in content_type.lower():
            return extract_text_from_docx(file_bytes)
        elif name.lower().endswith('.doc'):
            return extract_text_from_doc(file_bytes)
        else:
            return f"[Unsupported file type: {name}]"
    except Exception as e:
        return f"[Error extracting text from {name}: {str(e)}]"


def create_help_card():
    """Create a simple Adaptive Card with single Reformat CV button."""
    card = {
        "type": "AdaptiveCard",
        "$schema": "http://adaptivecards.io/schemas/adaptive-card.json",
        "version": "1.4",
        "body": [
            {
                "type": "TextBlock",
                "text": "Fernando Format",
                "weight": "Bolder",
                "size": "Medium"
            },
            {
                "type": "TextBlock",
                "text": "Upload a CV (PDF or Word) to reformat it with Meraki branding.",
                "wrap": True,
                "size": "Small",
                "color": "Accent"
            }
        ],
        "actions": [
            {
                "type": "Action.Submit",
                "title": "Reformat CV",
                "data": {"action": "reformat_cv"}
            }
        ]
    }

    return Attachment(
        content_type="application/vnd.microsoft.card.adaptive",
        content=card
    )


def create_start_new_card():
    """Create an Adaptive Card with just Start New button."""
    card = {
        "type": "AdaptiveCard",
        "$schema": "http://adaptivecards.io/schemas/adaptive-card.json",
        "version": "1.4",
        "body": [],
        "actions": [
            {
                "type": "Action.Submit",
                "title": "Start New",
                "data": {"action": "start_new"}
            }
        ]
    }

    return Attachment(
        content_type="application/vnd.microsoft.card.adaptive",
        content=card
    )


def generate_alternative_profile(cv_json: str) -> str:
    """Generate a short alternative candidate profile from CV data."""
    try:
        response = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "user", "content": ALTERNATIVE_PROFILE_PROMPT.format(cv_json=cv_json)}
            ],
            max_tokens=500,
            timeout=60
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print(f"[ERROR] Generating alternative profile: {e}")
        return ""


def is_reformat_trigger(text: str) -> bool:
    """Check if the message contains reformat trigger words."""
    if not text:
        return False
    text_lower = text.lower().strip()
    triggers = ["reformat", "format", "template", "/reformat", "!reformat"]
    words = text_lower.split()
    return any(trigger in words or text_lower.startswith(trigger) for trigger in triggers)


async def process_cv_reformat(cv_text: str, turn_context: TurnContext, show_start_new: bool = True, source_filename: str = None):
    """Process CV text and generate reformatted Word document with alternative profile."""
    try:
        # Step 1: Extract structured CV data using OpenAI
        response = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": CV_EXTRACTION_PROMPT},
                {"role": "user", "content": cv_text}
            ],
            max_tokens=8000,
            timeout=120
        )
        cv_json_text = response.choices[0].message.content

        # Step 2: Parse JSON and generate Word document
        cv_data = parse_cv_json(cv_json_text)
        doc_bytes = create_meraki_cv(cv_data)

        # Create filename from candidate name
        candidate_name = cv_data.get("name", "Candidate").replace(" ", "_")
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        filename = f"Meraki_CV_{candidate_name}_{timestamp}.docx"

        # Step 3: Generate alternative candidate profile
        alternative_profile = generate_alternative_profile(cv_json_text)

        # Step 4: Upload to Azure Blob Storage
        if blob_service_client:
            container_client = blob_service_client.get_container_client("cv-outputs")
            blob_client = container_client.get_blob_client(filename)
            blob_client.upload_blob(doc_bytes, overwrite=True)

            # Generate SAS URL valid for 7 days
            account_name = blob_service_client.account_name
            account_key = storage_connection_string.split("AccountKey=")[1].split(";")[0]
            sas_token = generate_blob_sas(
                account_name=account_name,
                container_name="cv-outputs",
                blob_name=filename,
                account_key=account_key,
                permission=BlobSasPermissions(read=True),
                expiry=datetime.utcnow() + timedelta(days=7)
            )
            download_url = f"https://{account_name}.blob.core.windows.net/cv-outputs/{filename}?{sas_token}"

            # Build response message
            response_text = (
                f"Here's the reformatted CV for **{cv_data.get('name', 'the candidate')}**:\n\n"
                f"[Download {filename}]({download_url})\n\n"
                f"_Link expires in 7 days_"
            )

            if alternative_profile:
                response_text += f"\n\n**Alternative Candidate Profile:**\n{alternative_profile}"

            if show_start_new:
                start_new_card = create_start_new_card()
                reply = Activity(
                    type="message",
                    text=response_text,
                    attachments=[start_new_card]
                )
                await turn_context.send_activity(reply)
            else:
                await turn_context.send_activity(response_text)
        else:
            await turn_context.send_activity("Error: Storage not configured for file uploads")

    except ValueError as e:
        error_msg = f"Error parsing CV data: {str(e)}"
        if source_filename:
            error_msg = f"Error processing **{source_filename}**: {str(e)}"
        await turn_context.send_activity(error_msg)
    except Exception as e:
        error_msg = f"Error generating Word document: {str(e)}"
        if source_filename:
            error_msg = f"Error processing **{source_filename}**: {str(e)}"
        await turn_context.send_activity(error_msg)


async def process_multiple_cvs(cv_files: list, turn_context: TurnContext):
    """Process multiple CV files separately."""
    total = len(cv_files)

    if total > 1:
        await turn_context.send_activity(f"Processing {total} CVs...")

    for i, (filename, cv_text) in enumerate(cv_files):
        is_last = (i == total - 1)

        if total > 1:
            await turn_context.send_activity(f"**Processing CV {i + 1} of {total}:** {filename}")

        show_button = is_last and total == 1
        await process_cv_reformat(cv_text, turn_context, show_start_new=show_button, source_filename=filename)

    if total > 1:
        start_new_card = create_start_new_card()
        reply = Activity(
            type="message",
            text=f"Finished processing {total} CVs.",
            attachments=[start_new_card]
        )
        await turn_context.send_activity(reply)


async def on_turn(turn_context: TurnContext):
    """Message handler for Fernando Format bot."""
    if turn_context.activity.type == "message":
        user_text = turn_context.activity.text or ""
        attachments = turn_context.activity.attachments or []
        conversation_id = turn_context.activity.conversation.id
        activity = turn_context.activity
        card_data = turn_context.activity.value

        # 1. Handle help commands
        if user_text.lower().strip() in ["help", "/help", "menu", "start", "hi", "hello", "hey"]:
            clear_pending_reformat(conversation_id)
            help_card = create_help_card()
            reply = Activity(type="message", attachments=[help_card])
            await turn_context.send_activity(reply)
            return

        # 2. Handle "Start New" button
        if card_data and card_data.get("action") == "start_new":
            clear_pending_reformat(conversation_id)
            help_card = create_help_card()
            reply = Activity(type="message", attachments=[help_card])
            await turn_context.send_activity(reply)
            return

        # Extract text from any file attachments
        cv_files = []
        extraction_errors = []
        for attachment in attachments:
            # Skip image attachments
            if attachment.content_type and attachment.content_type.startswith('image/'):
                continue
            # Handle HTML attachments - Teams sometimes sends CV content as HTML
            if attachment.content_type == 'text/html':
                html_content = attachment.content or ""
                if html_content:
                    extracted_text = extract_text_from_html(html_content)
                    if is_cv_content(extracted_text):
                        cv_files.append(("CV from Teams", extracted_text))
                        continue
                continue
            # Skip attachments with no download URL available
            has_download_url = (
                attachment.content_url or
                (isinstance(attachment.content, dict) and attachment.content.get('downloadUrl'))
            )
            if not has_download_url:
                continue
            name = attachment.name or "unknown"
            if not name and isinstance(attachment.content, dict):
                name = attachment.content.get("name", "unknown")
            extracted = await extract_text_from_attachment(attachment, turn_context)
            if extracted and not extracted.startswith('['):
                cv_files.append((name, extracted))
            elif extracted and extracted.startswith('['):
                extraction_errors.append(extracted)

        # If no files found via direct attachments, try Graph API to fetch from chat
        if len(cv_files) == 0 and len(attachments) > 0:
            token = get_graph_token()
            if token:
                user_aad_id = None
                if activity.from_property and hasattr(activity.from_property, 'aad_object_id'):
                    user_aad_id = activity.from_property.aad_object_id
                try:
                    graph_files = await get_files_from_chat(conversation_id, token, user_aad_id)
                    for gf in graph_files:
                        name = gf.get("name", "")
                        content_url = gf.get("content_url", "")
                        if name.lower().endswith(('.pdf', '.docx', '.doc')) and content_url:
                            try:
                                file_bytes = await download_file_from_sharepoint(content_url, token)
                                if name.lower().endswith('.pdf'):
                                    text = extract_text_from_pdf(file_bytes)
                                elif name.lower().endswith('.docx'):
                                    text = extract_text_from_docx(file_bytes)
                                elif name.lower().endswith('.doc'):
                                    text = extract_text_from_doc(file_bytes)
                                else:
                                    continue
                                if text and not text.startswith('['):
                                    cv_files.append((name, text))
                            except Exception as e:
                                print(f"[ERROR] Failed to download {name}: {e}")
                except Exception as e:
                    print(f"[ERROR] Graph API fetch failed: {e}")

        has_valid_files = len(cv_files) > 0
        has_text = bool(user_text.strip())

        # 3. Handle "Reformat CV" button press with no content
        if card_data and card_data.get("action") == "reformat_cv":
            if not has_valid_files and not has_text:
                set_pending_reformat(conversation_id)
                await turn_context.send_activity("Great! Send me the CV(s) - you can paste text or upload PDF/Word files.")
                return
            else:
                clear_pending_reformat(conversation_id)
                if has_valid_files:
                    await process_multiple_cvs(cv_files, turn_context)
                else:
                    await process_cv_reformat(user_text, turn_context)
                return

        # 4. Check for pending reformat state with content
        if get_pending_reformat(conversation_id) and (has_valid_files or has_text):
            clear_pending_reformat(conversation_id)
            if has_valid_files:
                await process_multiple_cvs(cv_files, turn_context)
            else:
                await process_cv_reformat(user_text, turn_context)
            return

        # 5. Check for reformat trigger words in message
        if is_reformat_trigger(user_text):
            if has_valid_files:
                await process_multiple_cvs(cv_files, turn_context)
                return
            elif has_text:
                content_to_process = user_text
                words = content_to_process.split(None, 1)
                if len(words) > 1 and is_reformat_trigger(words[0]):
                    content_to_process = words[1]
                if content_to_process.strip():
                    await process_cv_reformat(content_to_process, turn_context)
                    return
            set_pending_reformat(conversation_id)
            await turn_context.send_activity("Send me the CV(s) to reformat - paste text or upload PDF/Word files.")
            return

        # 6. File attached with no explicit trigger -> assume reformat
        if has_valid_files:
            await process_multiple_cvs(cv_files, turn_context)
            return

        # 7. Handle extraction errors
        if extraction_errors:
            for error in extraction_errors:
                await turn_context.send_activity(error)
            return

        # 8. Fallback -> show help card
        help_card = create_help_card()
        reply = Activity(type="message", attachments=[help_card])
        await turn_context.send_activity(reply)


_START_TIME = datetime.now().isoformat()

@app.route("/")
def home():
    return "Fernando Format bot is running!"

@app.route("/health")
def health():
    return {"status": "healthy", "service": "fernando-format", "started_at": _START_TIME}


@app.route("/api/messages", methods=["POST"])
def messages():
    if "application/json" in request.headers.get("Content-Type", ""):
        body = request.json
    else:
        return Response(status=415)

    activity = Activity().deserialize(body)
    auth_header = request.headers.get("Authorization", "")

    async def call_bot():
        await adapter.process_activity(activity, auth_header, on_turn)

    loop = asyncio.new_event_loop()
    asyncio.set_event_loop(loop)
    loop.run_until_complete(call_bot())
    loop.close()

    return Response(status=200)


if __name__ == "__main__":
    app.run()
